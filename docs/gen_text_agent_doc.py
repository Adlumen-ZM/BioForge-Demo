"""
生成 v0.1 Text Agent 技术方案（更新版）docx 文档。
运行方式：python3 Design_doc/gen_text_agent_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ─────────────────────────────────────────────
# 辅助函数
# ─────────────────────────────────────────────

def set_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def add_code_block(doc, code_lines):
    """添加代码块（灰底等宽字体段落）。"""
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        # 浅灰色背景
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)


def add_table_with_header(doc, headers, rows):
    """添加带表头的表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    # 数据行
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row):
            row_cells[c_idx].text = cell_text
    return table


# ─────────────────────────────────────────────
# 构建文档
# ─────────────────────────────────────────────

def build_doc():
    doc = Document()

    # ── 页面边距 ─────────────────────────────
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ══════════════════════════════════════════
    # 标题
    # ══════════════════════════════════════════
    title = doc.add_heading('v0.1 Text Agent 技术方案（更新版）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(doc, '本文档基于实际代码实现更新，反映 text_agent/ 模块的真实技术细节。')
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # 1. 模块概述
    # ══════════════════════════════════════════
    set_heading(doc, '1. 模块概述（技术需求分析）', 1)
    add_body(doc,
        'Text Agent 是 v0.1 版本中的核心信息抽取模块。给定一篇已选定的论文，'
        '将论文 PDF 传入 LLM，由 LLM 按照 field_dict_prompt.json 中定义的字段完成'
        '结构化抽取，并输出标准 JSON。本版本不引入 RAG，全文上下文一次性传入。'
    )
    add_body(doc,
        '与原始设计文档相比，本版本在以下方面进行了调整：'
    )
    add_bullet(doc, '使用 MiniMax-M2.5 模型（通过 OpenAI 兼容接口调用），而非 GPT-4o')
    add_bullet(doc, 'Reasoning 字段采用顶层平铺结构，而非嵌套 reasoning 对象')
    add_bullet(doc, 'ID 生成拆分为 4 个独立函数，paper_id 从数据库 MAX 值自增')
    add_bullet(doc, '新增 DOI 查重（阶段 8.5），防止同一论文重复写入')
    add_bullet(doc, '响应解析支持 4 种 JSON 提取策略，兼容 LLM 输出格式差异')

    # ══════════════════════════════════════════
    # 2. 技术目标
    # ══════════════════════════════════════════
    set_heading(doc, '2. 技术目标', 1)
    add_bullet(doc, 'JSON 格式合规、ID 格式合规、必填字段不为 null')
    add_bullet(doc, '核心溯源字段（text_to_* 三类）不为 null')
    add_bullet(doc, '信息抽取准确率达到 70%')
    add_bullet(doc, 'Trace 无论成功失败都写入，100% 可观测')
    add_bullet(doc, 'DOI 查重防止重复处理同一篇论文')

    # ══════════════════════════════════════════
    # 3. 输入规格
    # ══════════════════════════════════════════
    set_heading(doc, '3. 输入规格', 1)

    set_heading(doc, '3.1 输入规格', 2)
    add_bullet(doc, 'v0.1 由人工指定单篇论文，论文为 PDF 格式直接放在目录下的指定文件夹中')
    add_bullet(doc, '输入示例：v0.1_article_test.pdf（放于项目根目录）')

    set_heading(doc, '3.2 输入限制', 2)
    add_bullet(doc, 'v0.1 部分不特别处理图片、表格、附录文件等其它内容')
    add_bullet(doc, 'v0.1 部分不进行分块，全文一次性传入（MiniMax-M2.5 max_tokens=4096 输出）')
    add_bullet(doc, '不支持纯图像 PDF（依赖 PyMuPDF 的文字层提取）')

    # ══════════════════════════════════════════
    # 4. 方案设计
    # ══════════════════════════════════════════
    set_heading(doc, '4. 方案设计', 1)

    set_heading(doc, '4.1 方案概述', 2)
    add_body(doc,
        'v0.1 采用线性单轮流程，不引入 Agent 编排框架。整体由三个模块串联组成：'
        'Text Agent 负责调用 LLM 完成信息抽取，数据库模块负责结构化结果的持久化存储，'
        'Trace 模块负责记录每次执行的完整过程。'
    )
    add_body(doc, '系统使用两个独立的 SQLite 数据库文件，职责严格分离：')
    add_code_block(doc, [
        'hap_v01.db    业务数据库    存储抽取结果（paper_record + function_assay_evidence）',
        'hap_trace.db  Trace 数据库  存储执行过程（extraction_runs + trace_steps）',
    ])
    add_body(doc,
        '两个数据库通过 extraction_run_id 关联，业务数据可随时追溯至对应的 Trace 记录。'
    )

    set_heading(doc, '4.2 数据流', 2)
    add_body(doc, '实际实现的 TextAgent.run() 分为以下 10 个阶段：')

    add_body(doc, '阶段 1：PDF 文本提取')
    add_body(doc,
        '使用 PyMuPDF（fitz）从 PDF 逐页提取纯文本。过滤空白页，页间以双换行分隔。'
        '提取结果保留原始排版顺序，不做 OCR。'
    )
    add_code_block(doc, [
        'def extract_text_from_pdf(pdf_path: str) -> tuple[str | None, str | None]:',
        '    doc = fitz.open(pdf_path)',
        '    pages = []',
        '    for page_num, page in enumerate(doc, start=1):',
        '        text = page.get_text()',
        '        if text.strip():',
        '            pages.append(text)',
        '    doc.close()',
        '    full_text = "\\n\\n".join(pages)',
        '    return full_text, None',
    ])

    add_body(doc, '阶段 2：生成 paper_id')
    add_body(doc,
        'paper_id 从业务数据库查询当前最大值后自增生成（格式 P000001），'
        '不依赖外部序列机制。'
    )
    add_code_block(doc, [
        'def generate_paper_id() -> str:',
        '    # 查询 MAX(paper_id)，字典序与数值序一致（定长 6 位数字）',
        '    row = conn.execute("SELECT MAX(paper_id) FROM paper_record_v01_min").fetchone()',
        '    max_id = row[0]',
        '    if max_id is None:',
        '        return "P000001"',
        '    num = int(max_id[1:]) + 1',
        '    return f"P{num:06d}"',
    ])

    add_body(doc, '阶段 3：初始化 TraceLogger')
    add_body(doc,
        '生成 run_id（格式：run_YYYYMMDD_HHMMSS_xxxxxx），INSERT extraction_runs，'
        'run_status = "running"。从此阶段起所有失败均通过 TraceLogger 记录。'
    )

    add_body(doc, '阶段 4：构建 System/User Prompt')
    add_body(doc,
        'System Prompt 由 prompt_builder.build_system_prompt() 构建，'
        '从 field_dict_prompt.json 加载字段定义（详见第 5 节）。'
        'User Prompt 将论文全文嵌入，包含开始/结束标记。'
    )

    add_body(doc, '阶段 5：调用 LLM（MiniMax-M2.5）')
    add_body(doc,
        '通过 OpenAI 兼容接口调用 MiniMax-M2.5，API Key 从环境变量 MINIMAX_API_KEY 读取。'
        '在调用前后分别记录时间戳，收到响应后一次性写入 DB，避免干扰 response_time_ms 计算。'
    )
    add_code_block(doc, [
        'MODEL_NAME = "MiniMax-M2.5"',
        '',
        'client = OpenAI(',
        '    api_key=os.getenv("MINIMAX_API_KEY"),',
        '    base_url="https://api.minimaxi.com/v1",',
        ')',
        'response = client.chat.completions.create(',
        '    model=MODEL_NAME,',
        '    messages=[',
        '        {"role": "system", "content": system_prompt},',
        '        {"role": "user", "content": user_prompt},',
        '    ],',
        '    temperature=0.2,   # 低温度保证抽取结果的确定性',
        '    max_tokens=4096,',
        ')',
    ])

    add_body(doc, '阶段 6：INSERT trace_step')
    add_body(doc,
        '收到 LLM 响应后立即写入 trace_steps，step_status = "processing"。'
        '无论后续解析是否成功，原始响应都已持久化，可供故障恢复。'
    )

    add_body(doc, '阶段 7：解析 LLM 响应')
    add_body(doc,
        'parse_response() 先通过 4 种策略提取 JSON 文本，再执行 json.loads()，'
        '最后分离 reasoning 字段（详见第 5.3 节）。'
    )

    add_body(doc, '阶段 8.5：DOI 查重（新增）')
    add_body(doc,
        '解析出 DOI 后，查询业务数据库是否已存在相同 DOI 的记录。'
        '若已存在则跳过业务写入，以 success 状态结束 run，记录复用说明。'
    )
    add_code_block(doc, [
        'def find_paper_id_by_doi(doi: str) -> str | None:',
        '    row = conn.execute(',
        '        "SELECT paper_id FROM paper_record_v01_min WHERE doi = ?", (doi,)',
        '    ).fetchone()',
        '    return row[0] if row else None',
    ])

    add_body(doc, '阶段 9：业务字段校验')
    add_body(doc,
        '注入系统 ID（paper_id / record_id / fae_id）后，'
        '调用 BusinessDBWriter.validate_record() 执行枚举值、必填字段、ID 格式校验。'
    )

    add_body(doc, '阶段 10：写入业务数据库 + 完整性检查 + finalize')
    add_body(doc,
        '先写 paper_record，再批量写 FAE 子记录，最后执行 integrity_check（反向校验写入结果），'
        '最终调用 TraceLogger.finalize() 聚合 token 统计并写入最终状态。'
    )

    set_heading(doc, '4.3 关键设计决策', 2)

    add_body(doc, 'Trace 无论成功失败都写入')
    add_body(doc,
        'Trace 的首要价值在于故障复盘。阶段 6 收到 LLM 响应后立即写入原始响应，'
        '无论后续 JSON 解析是否成功，原始数据都已持久化。'
    )

    add_body(doc, 'v0.1 只做单次 LLM 调用')
    add_body(doc,
        'v0.1 的目标是验证 LLM 能否从论文中抽取结构化数据，单次调用是最直接的验证方式。'
        '多轮对话、自动修复、重试逻辑会引入额外变量，待 v0.2 再引入。'
    )

    add_body(doc, 'temperature=0.2 而非 0')
    add_body(doc,
        '低温度（0.2）在保证输出确定性的同时，保留一定的语言灵活性，'
        '避免强制 temperature=0 时某些模型产生重复 token 的问题。'
    )

    # ══════════════════════════════════════════
    # 5. Prompt 设计
    # ══════════════════════════════════════════
    set_heading(doc, '5. Prompt 设计', 1)

    set_heading(doc, '5.1 整体调用结构', 2)
    add_body(doc,
        'v0.1 采用单次 API 调用完成全部抽取。论文全文通过 User Prompt 传入，'
        'System Prompt 包含任务定义和字段字典，两者共同构成完整请求。'
    )
    add_code_block(doc, [
        'PROMPT_VERSION = "v0.1.0"   # 修改 System Prompt 内容时必须同步更新',
        'SCHEMA_VERSION = "v0.1"     # 修改 field_dict_prompt.json 时必须同步更新',
    ])

    set_heading(doc, '5.2 System Prompt', 2)
    add_body(doc,
        'System Prompt 由 build_system_prompt() 从 field_dict_prompt.json 动态构建，'
        '包含以下五个部分：'
    )
    add_bullet(doc, '任务描述：角色（生物医学文献挖掘助手）、研究背景（HAp 结合肽）、目标')
    add_bullet(doc, '输出格式规范：合法 JSON 对象，不使用 markdown 代码块，JSON 前后无多余文字')
    add_bullet(doc, '论文级字段定义（从 paper_level_fields 加载，含类型、是否必填、允许值、描述）')
    add_bullet(doc, 'FAE 子记录字段定义（从 fae_record_fields.fields 加载）')
    add_bullet(doc, 'Reasoning 字段说明 + 关键规则 + 输出示例（output_example）')
    add_code_block(doc, [
        'def build_system_prompt() -> str:',
        '    with open(_FIELD_DICT_PATH, encoding="utf-8") as f:',
        '        field_dict = json.load(f)',
        '    paper_fields = field_dict["paper_level_fields"]',
        '    fae_fields   = field_dict["fae_record_fields"]["fields"]',
        '    output_example = field_dict["output_example"]',
        '    # 构建字段说明文本，嵌入 prompt ...',
        '    return system_prompt',
    ])

    set_heading(doc, '5.3 User Prompt', 2)
    add_body(doc,
        'User Prompt 将论文全文以固定格式嵌入，包含开始/结束标记，指令简洁。'
    )
    add_code_block(doc, [
        'def build_user_prompt(paper_text: str) -> str:',
        '    return f"""请从以下科学论文全文中抽取 HAp 结合肽信息，',
        '按 System Prompt 中的规范输出一个合法的 JSON 对象。',
        '',
        '=== 论文全文开始 ===',
        '{paper_text}',
        '=== 论文全文结束 ===',
        '"""',
    ])

    set_heading(doc, '5.4 Reasoning 字段设计（与原设计文档的差异）', 2)
    add_body(doc,
        '原始设计文档中，reasoning 字段采用嵌套对象结构：'
    )
    add_code_block(doc, [
        '// 原始设计文档方案（已废弃）',
        '"reasoning": {',
        '    "entity_type_reason": "...",',
        '    "sequence_status_reason": "...",',
        '    "evidence_overall_level_reason": "...",',
        '    "trace_status_reason": "..."',
        '}',
    ])
    add_body(doc,
        '实际实现采用顶层平铺结构，reasoning 字段与业务字段同级，'
        '由 response_parser.py 在解析时识别并弹出，存入 trace_steps.llm_reasoning，'
        '不写入业务表：'
    )
    add_code_block(doc, [
        '// 实际实现方案（顶层平铺）',
        '"interaction_target": "enamel",',
        '"interaction_target_reason": "论文第2页 Introduction 明确写...",',
        '"summary_functions": ["adsorption", "remineralization"],',
        '"summary_functions_reason": "Abstract 写...",',
        '"evidence_overall_level": "in_vitro",',
        '"evidence_overall_level_reason": "全文实验使用 bovine enamel...",',
        '"trace_status": "complete",',
        '"trace_status_reason": "sequence、method、function、evidence 四类来源字段均有明确锚点"',
    ])
    add_body(doc, 'response_parser.py 中提取 reasoning 的逻辑：')
    add_code_block(doc, [
        '_REASONING_FIELDS = [',
        '    "interaction_target_reason",',
        '    "summary_functions_reason",',
        '    "evidence_overall_level_reason",',
        '    "trace_status_reason",',
        ']',
        '',
        'def parse_response(raw_response: str):',
        '    # ...提取 JSON 文本...',
        '    reasoning = {}',
        '    for field in _REASONING_FIELDS:',
        '        if field in data:',
        '            reasoning[field] = data.pop(field)  # 从业务数据中移除',
        '    llm_reasoning_json = json.dumps(reasoning, ensure_ascii=False) if reasoning else None',
        '    return data, llm_reasoning_json, None',
    ])

    set_heading(doc, '5.5 JSON 响应解析策略', 2)
    add_body(doc,
        'response_parser._extract_json_text() 按以下优先级依次尝试 4 种策略，'
        '兼容 LLM 输出格式差异：'
    )

    add_table_with_header(
        doc,
        ['策略', '触发条件', '处理方式'],
        [
            ['策略 1', 'LLM 输出以 { 开头', '直接使用原始文本（最常见情况）'],
            ['策略 2', '包含 ```json 代码块', '提取 ```json ... ``` 之间的内容'],
            ['策略 3', '包含普通 ``` 代码块', '提取 ``` ... ``` 之间的内容'],
            ['策略 4（兜底）', '均不匹配', '截取第一个 { 到最后一个 } 之间的内容'],
        ]
    )

    # ══════════════════════════════════════════
    # 6. ID 生成
    # ══════════════════════════════════════════
    set_heading(doc, '6. ID 生成', 1)
    add_body(doc,
        '原始设计文档中 ID 生成用单个 generate_ids() 函数处理。'
        '实际实现拆分为 4 个独立函数，职责更清晰：'
    )

    add_table_with_header(
        doc,
        ['函数', '职责', '格式示例'],
        [
            ['generate_paper_id()', '从 DB MAX(paper_id) 自增，无记录时从 P000001 开始', 'P000001'],
            ['find_paper_id_by_doi(doi)', 'DOI 查重，返回已存在的 paper_id 或 None', '—'],
            ['generate_record_id(paper_id, index=1)', '生成对象级记录编号，v0.1 index 固定为 1', 'P000001-R01'],
            ['generate_fae_id(record_id, index)', '生成 FAE 子记录编号，index 从 1 开始', 'P000001-R01-FAE01'],
        ]
    )

    add_body(doc,
        'TextAgent 中调用顺序：先 generate_paper_id() 生成 paper_id，'
        '解析出 LLM 结果后调用 find_paper_id_by_doi() 做查重，'
        '校验前注入 generate_record_id() 和 generate_fae_id()。'
    )
    add_code_block(doc, [
        '# TextAgent.run() 中 ID 注入逻辑',
        'record_id = generate_record_id(paper_id)          # P000001-R01',
        'parsed_output["paper_id"] = paper_id',
        'parsed_output["record_id"] = record_id',
        '',
        'fae_list = parsed_output.get("fae_records", [])',
        'for i, fae in enumerate(fae_list, start=1):',
        '    fae["fae_id"] = generate_fae_id(record_id, i)  # P000001-R01-FAE01',
    ])

    # ══════════════════════════════════════════
    # 7. 完整执行框架
    # ══════════════════════════════════════════
    set_heading(doc, '7. 完整执行框架', 1)
    add_code_block(doc, [
        '# text_agent/text_agent.py',
        'class TextAgent:',
        '    def __init__(self):',
        '        self._writer = BusinessDBWriter()  # 检查数据库表是否已就绪',
        '',
        '    def run(self, pdf_path: str) -> tuple[bool, str | None]:',
        '        # 阶段 1：PDF 文本提取',
        '        paper_text, err = extract_text_from_pdf(pdf_path)',
        '        if err: return False, f"[阶段1] {err}"',
        '',
        '        # 阶段 2/3：生成 paper_id',
        '        paper_id = generate_paper_id()',
        '',
        '        # 阶段 4：初始化 TraceLogger',
        '        logger = TraceLogger(paper_id=paper_id, model_name=MODEL_NAME,',
        '                             prompt_version=PROMPT_VERSION, schema_version=SCHEMA_VERSION)',
        '        run_id = logger.get_run_id()',
        '',
        '        # 阶段 5：构建 Prompt',
        '        system_prompt = build_system_prompt()',
        '        user_prompt   = build_user_prompt(paper_text)',
        '',
        '        # 阶段 6：调用 LLM',
        '        llm_result, err = call_llm(system_prompt, user_prompt)',
        '        if err:',
        '            logger.finalize("failed", error_message=err)',
        '            return False, err',
        '',
        '        # 阶段 7：INSERT trace_step（立即落地）',
        '        step_id, err = logger.insert_step(**llm_result_fields)',
        '',
        '        # 阶段 8：解析响应',
        '        parsed_output, llm_reasoning, err = parse_response(llm_result["raw_response"])',
        '        if err:',
        '            logger.update_step(step_id, step_status="parse_error", error_detail=err)',
        '            logger.finalize("failed", error_message=err)',
        '            return False, err',
        '',
        '        # 阶段 8.5：DOI 查重',
        '        doi = parsed_output.get("doi")',
        '        if doi and find_paper_id_by_doi(doi):',
        '            logger.finalize("success", error_message="DOI 已存在，跳过写入")',
        '            return True, "已存在"',
        '',
        '        # 阶段 9：注入 ID + 字段校验',
        '        parsed_output["paper_id"] = paper_id',
        '        parsed_output["record_id"] = generate_record_id(paper_id)',
        '        for i, fae in enumerate(fae_list, 1):',
        '            fae["fae_id"] = generate_fae_id(record_id, i)',
        '        ok, err = self._writer.validate_record(parsed_output)',
        '',
        '        # 阶段 10：写入业务 DB + 完整性检查 + finalize',
        '        self._writer.write_paper_record(parsed_output, run_id)',
        '        self._writer.write_fae_records(fae_list, record_id)',
        '        check_status, check_detail = self._writer.integrity_check(paper_id, len(fae_list))',
        '        logger.finalize("success" if check_status == "passed" else "failed")',
        '        return check_status == "passed", check_detail',
    ])

    # ══════════════════════════════════════════
    # 8. 错误处理
    # ══════════════════════════════════════════
    set_heading(doc, '8. 错误处理', 1)
    add_body(doc,
        'Text Agent 在执行过程中可能遇到以下错误类型。所有错误均写入 Trace，'
        'Trace 写入本身不受业务错误影响。'
    )

    add_table_with_header(
        doc,
        ['错误类型', '触发场景', '处理策略', 'step_status'],
        [
            ['PDF 文本提取失败', '文件不存在、无文字层、损坏',
             '直接返回 False，未初始化 TraceLogger 故无法写 Trace', 'run_status=failed（阶段1）'],
            ['API Key 未设置', 'MINIMAX_API_KEY 环境变量为空',
             '返回错误描述，finalize failed', 'api_error'],
            ['API 速率限制', 'HTTP 429',
             '返回错误描述，finalize failed', 'api_error'],
            ['API 服务器错误', 'HTTP 500',
             '返回错误描述，finalize failed', 'api_error'],
            ['API 调用超时', 'timeout 关键字', '返回错误描述，finalize failed', 'timeout'],
            ['API Key 无效', 'HTTP 401 / unauthorized', '返回错误描述，finalize failed', 'api_error'],
            ['JSON 解析失败', 'LLM 输出非合法 JSON（4 种策略均失败）',
             '记录原始响应，finalize failed', 'parse_error'],
            ['字段校验失败', '枚举值超出范围或必填字段为 null',
             '记录错误字段，finalize failed', 'schema_error'],
            ['DOI 重复', '同一 DOI 已在数据库中',
             '跳过业务写入，以 success 结束，记录复用说明', 'success'],
            ['完整性检查失败', 'FAE 记录数不符或核心字段为 null',
             '数据已写入但标记 failed，记录差异详情', 'run_status=failed'],
        ]
    )

    # ══════════════════════════════════════════
    # 9. 模块结构
    # ══════════════════════════════════════════
    set_heading(doc, '9. 模块结构', 1)

    add_table_with_header(
        doc,
        ['文件', '职责'],
        [
            ['text_agent.py', 'TextAgent 主类，10 阶段流程编排'],
            ['pdf_extractor.py', 'PyMuPDF PDF 文本提取（逐页、过滤空白页）'],
            ['prompt_builder.py', 'System/User Prompt 构建，加载 field_dict_prompt.json'],
            ['llm_client.py', 'MiniMax-M2.5 API 调用（OpenAI 兼容接口）'],
            ['response_parser.py', 'LLM 响应解析：4 策略 JSON 提取 + reasoning 字段分离'],
            ['id_generator.py', 'paper_id 自增生成、DOI 查重、record_id/fae_id 生成'],
            ['field_dict_prompt.json', '嵌入 System Prompt 的字段字典（v0.1 版本）'],
        ]
    )

    # ══════════════════════════════════════════
    # 10. 已知局限
    # ══════════════════════════════════════════
    set_heading(doc, '10. 已知局限', 1)

    add_table_with_header(
        doc,
        ['局限', '原因', '计划版本'],
        [
            ['不支持扫描版 PDF', 'PyMuPDF 依赖 PDF 文字层', 'v0.2 引入 OCR 方案'],
            ['输出 token 上限 4096', 'max_tokens=4096，FAE 记录较多时可能截断', 'v0.2 调整'],
            ['不支持批量处理', 'v0.1 单篇验证，无任务队列', 'v0.3 引入批量处理流程'],
            ['不支持图表信息抽取', 'PyMuPDF 纯文本提取不含图片内容', 'v0.2 引入 Figure/Table Agent'],
            ['不做自动修复', 'v0.1 验证基础能力，避免引入额外变量', 'v0.2 视验证结果决定是否引入'],
            ['单模型', 'v0.1 使用 MiniMax-M2.5 单模型验证', 'v0.2 引入多模型适配层'],
            ['paper_id 生成存在竞态窗口', 'MAX 自增在并发场景下不安全', 'v0.2 并发处理时改为行锁或序列机制'],
        ]
    )

    return doc


if __name__ == '__main__':
    import os
    output_path = os.path.join(os.path.dirname(__file__), 'v0.1 Text Agent 技术方案（更新版）.docx')
    doc = build_doc()
    doc.save(output_path)
    print(f'文档已生成：{output_path}')
